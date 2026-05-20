"""简历解析器单元测试 — 纯解析逻辑,不依赖 DB / Celery / 存储。

覆盖 M5 修复点:
- DOCX 表格 / 嵌套表格 / 页眉页脚的 email/phone 抽取
- PDF 双栏被打散成"一行一字"时,联系信息仍能从压缩文本里恢复
- _maybe_dedup_garbled 对 ≥30% 短行的输入做折叠

注意:PDF 真实样例(双栏 / 水印 / OCR 残文)很难纯代码构造,我们用
模拟过的"已经被 PDF 引擎抽出来的乱序文本"直接喂 _extract_email_phone
和 _maybe_dedup_garbled,验证关键链路。完整 PDF 端到端覆盖留给手工测试。
"""
from __future__ import annotations

import io

import pytest
from docx import Document

from app.services.resume_parser import (
    _extract_email_phone,
    _maybe_dedup_garbled,
    parse_resume,
)

DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


# ---------------------------------------------------------------------------
# _extract_email_phone:乱序场景兜底
# ---------------------------------------------------------------------------


def test_extract_normal_text():
    """正常文本走原文匹配路径。"""
    text = "联系方式:邮箱 li@bar.com 手机 13900001111"
    e, p = _extract_email_phone(text)
    assert e == "li@bar.com"
    assert p == "13900001111"


def test_extract_garbled_two_column_pdf():
    """模拟 PDF 双栏被打散:每个字符独占一行,中文标签做边界。"""
    src = "姓名: 张三  邮箱: zhangsan@example.com 手机: 13812345678"
    garbled = "\n".join(list(src))  # 每个字符一行
    e, p = _extract_email_phone(garbled)
    # 原文里每行 1 字符,正则匹配不到完整邮箱;走压缩兜底
    assert e == "zhangsan@example.com"
    assert p == "13812345678"


def test_extract_garbled_with_watermark():
    """模拟带随机水印 + 乱序的 PDF 文本(用户实际反馈的样例缩样)。"""
    sample = """4
d
c
b
e
rJhMWPvmgeOWduvU8uYwZtVE0S90_dX1edea4dcbe4b930a1
Jr Jr
M M
项目: Spring Boot + Spring Cloud
邮箱: candidate@163.com 手机: 13700009999
"""
    e, p = _extract_email_phone(sample)
    assert e == "candidate@163.com"
    assert p == "13700009999"


def test_extract_no_match_returns_none():
    e, p = _extract_email_phone("纯文本,无联系方式。")
    assert e is None
    assert p is None


# ---------------------------------------------------------------------------
# _maybe_dedup_garbled:UI 展示用的清洗
# ---------------------------------------------------------------------------


def test_dedup_garbled_folds_short_runs():
    """连续 ≥5 行短行被折叠为单行。"""
    text = "\n".join(["a", "b", "c", "d", "e", "f", "正常一行", "x"])
    cleaned = _maybe_dedup_garbled(text)
    assert "abcdef" in cleaned
    assert "正常一行" in cleaned


def test_dedup_garbled_keeps_normal_text():
    """短行占比 < 30% 时,原样返回。"""
    text = "\n".join(["normal line " + str(i) for i in range(10)] + ["a"])
    cleaned = _maybe_dedup_garbled(text)
    assert cleaned == text


def test_dedup_garbled_short_isolated_lines_kept():
    """只有 1-2 个短行夹在长行之间时,不要误折叠。"""
    text = "Spring Boot 项目\n-\nMyBatis 模块"
    cleaned = _maybe_dedup_garbled(text)
    # 短行占比是 1/3 ≈ 33%,触发折叠;但连续短行只有 1 行 < 5,所以不合并
    assert "Spring Boot 项目" in cleaned
    assert "MyBatis 模块" in cleaned


# ---------------------------------------------------------------------------
# DOCX:表格 / 页眉页脚 / 嵌套
# ---------------------------------------------------------------------------


def _docx_bytes(populate) -> bytes:
    doc = Document()
    populate(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_paragraph_only():
    """旧路径(只有段落)仍能抽到联系方式。"""

    def fill(doc):
        doc.add_paragraph("姓名: 王五")
        doc.add_paragraph("邮箱: w5@x.cn  手机: 13988887777")

    r = parse_resume("p.docx", DOCX_MIME, _docx_bytes(fill))
    assert r.email == "w5@x.cn"
    assert r.phone == "13988887777"


def test_docx_contact_in_table():
    """联系方式放在 2x2 表格里(简历模板常见形式)。"""

    def fill(doc):
        doc.add_paragraph("候选人简历")
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "邮箱"
        t.cell(0, 1).text = "tabuser@company.com"
        t.cell(1, 0).text = "手机"
        t.cell(1, 1).text = "13611112222"

    r = parse_resume("t.docx", DOCX_MIME, _docx_bytes(fill))
    assert r.email == "tabuser@company.com"
    assert r.phone == "13611112222"
    # 表格单元格的文本必须出现在 raw_text 里
    assert "tabuser@company.com" in r.raw_text


def test_docx_contact_in_header():
    """联系方式放在页眉里(高级简历模板把姓名/电话顶在页眉)。"""

    def fill(doc):
        doc.add_paragraph("候选人简历正文")
        section = doc.sections[0]
        section.header.paragraphs[0].text = (
            "hdr@inheader.com 13522223333"
        )

    r = parse_resume("hdr.docx", DOCX_MIME, _docx_bytes(fill))
    assert r.email == "hdr@inheader.com"
    assert r.phone == "13522223333"


def test_docx_contact_in_nested_table():
    """嵌套表格(单元格里再放表格)同样能抽到。"""

    def fill(doc):
        outer = doc.add_table(rows=1, cols=1)
        inner = outer.cell(0, 0).add_table(rows=1, cols=2)
        inner.cell(0, 0).text = "邮箱"
        inner.cell(0, 1).text = "nested@deep.com 13877776666"

    r = parse_resume("nested.docx", DOCX_MIME, _docx_bytes(fill))
    assert r.email == "nested@deep.com"
    assert r.phone == "13877776666"


def test_docx_drawingml_textbox_with_alternate_content():
    """模板把内容塞进 ``mc:AlternateContent`` 包裹的 DrawingML 文本框时,
    旧实现会:
    - 漏抽:python-docx ``Document.paragraphs`` 不递归进 ``<w:drawing>``
    - 重复:同一段在 ``<mc:Choice>`` 与 ``<mc:Fallback>`` 都有,iter() 各拿一遍

    新实现走 raw zip + lxml,只走 Choice 路径,本测试用最小可复现样本验证。
    """
    import zipfile

    # 直接构造一份最小 docx — python-docx 写不出 mc:AlternateContent
    document_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
  <w:body>
    <w:p><w:r><w:t>正文标题</w:t></w:r></w:p>
    <w:p>
      <w:r>
        <mc:AlternateContent>
          <mc:Choice Requires="wps">
            <w:drawing>
              <inner>
                <w:txbxContent>
                  <w:p><w:r><w:t>姓名: 王二麻子</w:t></w:r></w:p>
                  <w:p><w:r><w:t>邮箱: wang@test.cn 手机: 13988887777</w:t></w:r></w:p>
                </w:txbxContent>
              </inner>
            </w:drawing>
          </mc:Choice>
          <mc:Fallback>
            <inner>
              <w:p><w:r><w:t>姓名: 王二麻子</w:t></w:r></w:p>
              <w:p><w:r><w:t>邮箱: wang@test.cn 手机: 13988887777</w:t></w:r></w:p>
            </inner>
          </mc:Fallback>
        </mc:AlternateContent>
      </w:r>
    </w:p>
  </w:body>
</w:document>
""".encode()

    content_types = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""

    rels = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)

    r = parse_resume("textbox.docx", DOCX_MIME, buf.getvalue())

    # 1) email/phone 抽到
    assert r.email == "wang@test.cn"
    assert r.phone == "13988887777"

    # 2) 不重复:含"姓名"那行只出现一次
    assert r.raw_text.count("姓名: 王二麻子") == 1
    assert r.raw_text.count("邮箱: wang@test.cn") == 1

    # 3) 正文与文本框都覆盖
    assert "正文标题" in r.raw_text


# ---------------------------------------------------------------------------
# parse_resume / parse_quick 顶层 API
# ---------------------------------------------------------------------------


def test_parse_resume_txt_basic():
    data = "邮箱: test@a.com 电话: 13511112222 技能: Python Java".encode()
    r = parse_resume("候选人简历.txt", "text/plain", data)
    assert r.email == "test@a.com"
    assert r.phone == "13511112222"
    # 技能词典命中
    assert "Python" in r.skills
    assert "Java" in r.skills


def test_parse_resume_returns_partial_on_no_contact():
    """无联系方式时不抛错,字段为 None。"""
    r = parse_resume("note.txt", "text/plain", b"some random text without contact")
    assert r.email is None
    assert r.phone is None
    # raw_text 仍可用
    assert "random text" in r.raw_text


@pytest.mark.parametrize(
    "fn,expected",
    [
        ("张三-简历.pdf", "张三"),
        ("CV_李四_2024.pdf", "李四 2024"),
        ("resume.pdf", None),  # 全是噪声词
    ],
)
def test_name_hint_from_filename(fn, expected):
    from app.services.resume_parser import _name_hint_from_filename

    assert _name_hint_from_filename(fn) == expected
